#!/usr/bin/env python3
"""Generate a DOCX report for the Single-Cell Intelligence Agent.

Produces a Word document summarizing the agent's capabilities, seed
data statistics, and collection configuration.

Usage:
    python scripts/generate_docx.py [--output report.docx]
"""

import argparse
import logging
import sys
from pathlib import Path

# Ensure project root on sys.path
PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger(__name__)


def generate_report(output_path: str = "single_cell_agent_report.docx"):
    """Generate a DOCX report with agent summary information."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError:
        logger.error(
            "python-docx not installed. Install with: pip install python-docx"
        )
        # Fallback: generate a text file
        logger.info("Generating text report instead ...")
        _generate_text_report(output_path.replace(".docx", ".txt"))
        return

    from src.ingest.cellxgene_parser import get_cell_type_count, get_cell_lineages
    from src.ingest.marker_parser import get_marker_count, get_marker_sources
    from src.ingest.tme_parser import get_tme_profile_count, get_cancer_types
    from src.models import SCWorkflowType

    doc = Document()
    doc.add_heading("Single-Cell Intelligence Agent Report", level=0)
    doc.add_paragraph(
        "This report summarizes the capabilities and seed data for the "
        "HCLS AI Factory Single-Cell Intelligence Agent."
    )

    # Seed Data Statistics
    doc.add_heading("Seed Data Statistics", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Data Source"
    table.cell(0, 1).text = "Count"
    table.cell(1, 0).text = "Cell Type Records (CellxGene/HCA)"
    table.cell(1, 1).text = str(get_cell_type_count())
    table.cell(2, 0).text = "Marker Gene Records (CellMarker/PanglaoDB)"
    table.cell(2, 1).text = str(get_marker_count())
    table.cell(3, 0).text = "TME Profiles (Cancer Atlas)"
    table.cell(3, 1).text = str(get_tme_profile_count())

    # Workflows
    doc.add_heading("Analysis Workflows", level=1)
    for wf in SCWorkflowType:
        doc.add_paragraph(wf.value.replace("_", " ").title(), style="List Bullet")

    # Cell Lineages
    doc.add_heading("Cell Lineages Covered", level=1)
    for lineage in get_cell_lineages():
        doc.add_paragraph(lineage.replace("_", " ").title(), style="List Bullet")

    # Cancer Types
    doc.add_heading("Cancer TME Profiles", level=1)
    for cancer in get_cancer_types():
        doc.add_paragraph(cancer, style="List Bullet")

    doc.save(output_path)
    logger.info("DOCX report generated: %s", output_path)


def _generate_text_report(output_path: str):
    """Fallback text report when python-docx is not available."""
    from src.ingest.cellxgene_parser import get_cell_type_count
    from src.ingest.marker_parser import get_marker_count
    from src.ingest.tme_parser import get_tme_profile_count
    from src.models import SCWorkflowType

    lines = [
        "Single-Cell Intelligence Agent Report",
        "=" * 40,
        "",
        "Seed Data Statistics:",
        f"  Cell Type Records: {get_cell_type_count()}",
        f"  Marker Gene Records: {get_marker_count()}",
        f"  TME Profiles: {get_tme_profile_count()}",
        "",
        "Analysis Workflows:",
    ]
    for wf in SCWorkflowType:
        lines.append(f"  - {wf.value.replace('_', ' ').title()}")

    Path(output_path).write_text("\n".join(lines))
    logger.info("Text report generated: %s", output_path)


def main():
    """Parse arguments and generate report."""
    parser = argparse.ArgumentParser(description="Generate single-cell agent DOCX report")
    parser.add_argument(
        "--output",
        default="single_cell_agent_report.docx",
        help="Output file path (default: single_cell_agent_report.docx)",
    )
    args = parser.parse_args()
    generate_report(args.output)


if __name__ == "__main__":
    main()
